import threading
from pathlib import Path


from docx import Document
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font

from src.doc_agent.doc_tools._helpers import sandbox_root, ensure_project_path, generate_agent_tools
from src.doc_agent.doc_tools.registry import tool_registry


class OfficeFileIOToolkit:
    """
    office操作工具集
    """
    # 新增全局互斥锁
    _office_lock = threading.Lock()

    def __init__(self, custom_sandbox: Path | None = None):
        # 不传参自动读取本地默认路径，兼容原有写法
        if custom_sandbox is not None:
            self.sandbox_root = custom_sandbox.resolve()
            self.sandbox_root.mkdir(parents=True, exist_ok=True)
        else:
            self.sandbox_root = sandbox_root

    # -------------------------- 公共校验 --------------------------

    def _safe_file_path(self,raw_path: str) -> Path | str:
        try:
            return ensure_project_path(raw_path, self.sandbox_root)
        except ValueError as e:
            return f"[ERROR] {e}"


    # -------------------------- Word 读写工具 --------------------------


    def read_docx(self,file_path: str) -> str:
        """读取docx文档，提取全部纯文本，沙盒路径校验"""
        path = self._safe_file_path(file_path)
        if isinstance(path, str):
            return path

        if not path.exists():
            return "[ERROR] 文件不存在"
        if path.suffix.lower() != ".docx":
            return "[ERROR] 仅支持 .docx 格式，不支持旧版 .doc"

        try:
            # 创建Document对象
            doc = Document(path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            return "\n".join(full_text)
        except Exception as e:
            return f"[ERROR] docx解析失败：{str(e)}"

    def write_docx(self, file_path: str, content: str) -> str:
        """新建/覆盖写入docx文档，写入多行正文"""
        path = self._safe_file_path(file_path)
        if isinstance(path, str):
            return path

        path.parent.mkdir(parents=True, exist_ok=True)

        with self._office_lock:
            if path.exists():
                try:
                    old_doc = Document(path)
                    old_lines = []
                    for para in old_doc.paragraphs:
                        if para.text.strip():
                            old_lines.append(para.text.strip())
                    old_text = "\n".join(old_lines)
                    new_lines = content.splitlines()
                    new_text = "\n".join([line.strip() for line in new_lines if line.strip()])
                    if old_text == new_text:
                        return f"[INFO]文档内容无变更，跳过保存: {path}"
                except Exception:
                    pass

            doc = Document()
            lines = content.splitlines()
            for line in lines:
                doc.add_paragraph(line)

            try:
                doc.save(path)
                return f"成功生成docx文件：{path}"
            except Exception as e:
                return f"[ERROR] 写入docx失败：{str(e)}"

    # -------------------------- Excel 读写工具 --------------------------


    def read_xlsx(self, file_path: str) -> str:
        """读取xlsx表格，转为纯文本"""
        path = self._safe_file_path(file_path)
        if isinstance(path, str):
            return path

        if not path.exists():
            return "[ERROR] 文件不存在"
        if path.suffix.lower() != ".xlsx":
            return "[ERROR] 仅支持 .xlsx"

        try:
            wb = load_workbook(path, read_only=True)
            result = []
            for sheet in wb.worksheets:
                result.append(f"===== 工作表：{sheet.title} =====")
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    result.append(row_text)
            # 返回结果
            return "\n".join(result)
        except Exception as e:
            return f"[ERROR] xlsx解析失败：{str(e)}"

    def write_xlsx(self, file_path: str, data: list, sheet_name: str = "Sheet1", has_header: bool = True) -> str:
        """
        写入Excel文件，支持单工作表，可设置表头、自动适配列宽
        :param file_path: 沙盒内保存路径，如 "销售数据/7月销售统计表.xlsx"
        :param data: 二维列表格式的表格数据，如 [["姓名", "销售额"], ["张三", 10000], ["李四", 20000]]
        :param sheet_name: 工作表名称，默认Sheet1
        :param has_header: 首行是否为表头，默认True，表头会自动加粗
        """
        path = self._safe_file_path(file_path)
        if isinstance(path, str):
            return path

        if not data or not isinstance(data, list) or len(data) == 0:
            return "[ERROR] 表格数据不能为空，必须为二维列表格式"
        if not isinstance(data[0], list):
            return "[ERROR] 表格数据必须为二维列表，每一行是一个子列表"
        if path.suffix.lower() != ".xlsx":
            return "[ERROR] 仅支持 .xlsx 格式写入"

        path.parent.mkdir(parents=True, exist_ok=True)

        with self._office_lock:
            if path.exists():
                try:
                    wb_old = load_workbook(path, read_only=True)
                    sheet = wb_old[sheet_name]
                    old_table = []
                    for row in sheet.iter_rows(values_only=True):
                        row_list = [cell if cell is not None else "" for cell in row]
                        old_table.append(row_list)
                    wb_old.close()
                    if old_table == data:
                        return f"[INFO]表格数据无变更，跳过保存: {path}"
                except Exception:
                    pass

            try:
                wb = Workbook()
                ws = wb.active
                ws.title = sheet_name

                for row_idx, row_data in enumerate(data):
                    for col_idx, cell_value in enumerate(row_data):
                        ws.cell(row=row_idx + 1, column=col_idx + 1, value=cell_value if cell_value is not None else "")

                if has_header:
                    header_font = Font(bold=True)
                    for col in range(1, len(data[0]) + 1):
                        ws.cell(row=1, column=col).font = header_font

                for col in ws.columns:
                    max_length = 0
                    column = col[0].column_letter
                    for cell in col:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except (TypeError, ValueError):
                            pass
                    adjusted_width = (max_length + 2) * 1.2
                    ws.column_dimensions[column].width = adjusted_width

                wb.save(path)
                return f"成功生成Excel文件：{path}，共写入{len(data)}行数据，工作表：{sheet_name}"
            except Exception as e:
                return f"[ERROR] 写入xlsx失败：{str(e)}"


# 全局唯一实例
office_tools = OfficeFileIOToolkit()

# 注册工具列表
tools_list = generate_agent_tools(office_tools, skip_names=["generate_tool_list"])

# 注册工具
tool_registry.register_many(tools_list)
